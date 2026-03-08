
import streamlit as st
import os
import sqlite3
import logging
import pandas as pd
import json
import time
from io import BytesIO
from docx import Document
from src.generators.llm_client import LLMClient, PROVIDER_DEFAULTS
from src.generators.prompts import QUESTION_SYSTEM_MSG
import config

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(name)s: %(message)s",
)
logger = logging.getLogger(__name__)

DB_NAME = config.DB_NAME

# --- CONFIGURATION ---
st.set_page_config(
    page_title="Maarif-Gen",
    page_icon="🎓",
    layout="wide",
    initial_sidebar_state="expanded"
)

# --- DATABASE HELPERS ---
def get_db_connection():
    conn = sqlite3.connect(DB_NAME)
    conn.row_factory = sqlite3.Row
    return conn

@st.cache_data
def get_subjects():
    conn = get_db_connection()
    subjects = [dict(row) for row in conn.execute("SELECT * FROM subjects ORDER BY name").fetchall()]
    conn.close()
    return subjects

@st.cache_data
def get_grades(subject_id):
    conn = get_db_connection()
    grades = [dict(row) for row in conn.execute("SELECT * FROM grades WHERE subject_id = ? ORDER BY CAST(REPLACE(level_name, '. Sınıf', '') AS INTEGER)", (subject_id,)).fetchall()]
    conn.close()
    return grades

@st.cache_data
def get_units(grade_id):
    conn = get_db_connection()
    units = [dict(row) for row in conn.execute("SELECT * FROM units WHERE grade_id = ? ORDER BY unit_no", (grade_id,)).fetchall()]
    conn.close()
    return units

@st.cache_data
def get_outcomes(unit_id):
    conn = get_db_connection()
    outcomes = [dict(row) for row in conn.execute("SELECT * FROM outcomes WHERE unit_id = ? ORDER BY code", (unit_id,)).fetchall()]
    conn.close()
    return outcomes

@st.cache_data
def get_components(outcome_id):
    conn = get_db_connection()
    comps = [dict(row) for row in conn.execute("SELECT * FROM components WHERE outcome_id = ? ORDER BY code", (outcome_id,)).fetchall()]
    conn.close()
    return comps

# --- DOCX GENERATOR ---
def append_to_subject_docx(subject_name, outcome, components, questions):
    """Appends generated questions to a persistent Word document named after the subject."""
    os.makedirs("SoruBankasi", exist_ok=True)
    filepath = os.path.join("SoruBankasi", f"{subject_name}.docx")
    
    # Dosya zaten varsa aç, yoksa yeni oluştur
    if os.path.exists(filepath):
        doc = Document(filepath)
    else:
        doc = Document()
        doc.add_heading(f'Maarif-Gen {subject_name} Soru Bankası', 0)
    
    outcome_label = f"{outcome['code']} - {outcome.get('description', '')}"
    doc.add_heading(outcome_label, level=1)
    
    # Context
    if st.session_state.generated_context:
        doc.add_heading("Bağlam", level=2)
        doc.add_paragraph(st.session_state.generated_context)

    # Questions
    # Iterate over outcome components to preserve order and get texts
    added_any = False
    for comp in components:
        comp_code = comp['code']
        q_data = questions.get(comp_code)

        # Geçilen veya üretilmemiş bileşenleri atla
        if not q_data or q_data.get('_skipped'):
            continue

        added_any = True
        doc.add_heading(f"Süreç Bileşeni: {comp_code}) {comp.get('description', '')}", level=2)

        # Question Text
        q_text = q_data.get('question_text', '')
        doc.add_paragraph(q_text)

        # Rubric
        doc.add_heading('Rubrik', level=3)
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Kriter'
        hdr_cells[1].text = 'Puan'

        for item in q_data.get('rubric', []):
            row_cells = table.add_row().cells
            row_cells[0].text = str(item.get('criteria', ''))
            row_cells[1].text = str(item.get('points', ''))

        # Doğru Cevap Özeti
        doc.add_heading('Doğru Cevap Özeti', level=3)
        doc.add_paragraph(str(q_data.get('correct_answer_summary', '')))

        doc.add_page_break()
             
    if added_any:
        doc.save(filepath)
    
    return filepath

def append_to_finetuning_dataset(client, subject_name, outcome, components, questions, generated_context, impl_guide, rubric_points):
    """Saves the prompt and result in JSONL format for future LLM fine-tuning."""
    os.makedirs("SoruBankasi", exist_ok=True)
    filepath = os.path.join("SoruBankasi", "finetune_dataset.jsonl")
    
    with open(filepath, "a", encoding="utf-8") as f:
        for comp in components:
            comp_code = comp['code']
            q_data = questions.get(comp_code)

            # Geçilen veya üretilmemiş bileşenleri atla
            if not q_data or q_data.get('_skipped'):
                continue

            ctx = _build_context_prompt(generated_context, impl_guide, rubric_points)

            outcome_desc = outcome.get('description', '') or outcome['code']
            outcome_dict = {"code": outcome['code'], "text": f"{comp['description']} ({outcome_desc})"}
            user_prompt = client._create_prompt(outcome_dict, ctx)

            # _elapsed ve _skipped gibi dahili alanları temizle
            clean_q = {k: v for k, v in q_data.items() if not k.startswith('_')}
            assistant_response = json.dumps(clean_q, ensure_ascii=False)
            
            record = {
                "messages": [
                    {"role": "system", "content": QUESTION_SYSTEM_MSG},
                    {"role": "user", "content": user_prompt},
                    {"role": "assistant", "content": assistant_response}
                ],
                "metadata": {
                    "subject": subject_name,
                    "code": outcome['code'],
                    "component": comp_code
                }
            }
            f.write(json.dumps(record, ensure_ascii=False) + "\n")


def _text_height(text: str, min_height: int = 150) -> int:
    """İçeriğe göre text_area yüksekliğini hesaplar."""
    if not text:
        return min_height
    lines = text.split('\n')
    # Uygulama "wide" (geniş) tasarıma sahip olduğu için tek satıra 75 yerine ortalama 130 karakter sığar.
    total_lines = sum(max(1, -(-len(line) // 130)) for line in lines)
    return max(min_height, (total_lines * 26) + 30)


def _build_context_prompt(generated_context: str, impl_guide: str, rubric_points: int) -> str:
    """Assembles the context string passed to LLM question generation calls."""
    ctx = f"Bağlam:\n{generated_context}\n\n"
    if impl_guide:
        ctx += f"Uygulama Esasları: {impl_guide}\n\n"
    ctx += f"ÖNEMLİ: Rubrik toplam {rubric_points} puan olacak şekilde ayarlanmalıdır."
    return ctx


# --- SESSION STATE ---
if 'selected_outcome' not in st.session_state:
    st.session_state.selected_outcome = None
if 'generated_context' not in st.session_state:
    st.session_state.generated_context = ""
if 'context_elapsed' not in st.session_state:
    st.session_state.context_elapsed = None
if 'generated_questions' not in st.session_state:
    st.session_state.generated_questions = {}
if 'last_outcome_code' not in st.session_state:
    st.session_state.last_outcome_code = None
if 'w_step' not in st.session_state:
    st.session_state.w_step = 1
if 's_subj' not in st.session_state:
    st.session_state.s_subj = None
if 's_grade' not in st.session_state:
    st.session_state.s_grade = None
if 's_unit' not in st.session_state:
    st.session_state.s_unit = None
if 'llm_provider' not in st.session_state:
    st.session_state.llm_provider = "local"
if 'llm_api_key' not in st.session_state:
    st.session_state.llm_api_key = ""

# --- LLM CLIENT ---
@st.cache_resource
def get_llm_client(provider: str, api_key: str, model: str):
    defaults = PROVIDER_DEFAULTS.get(provider, PROVIDER_DEFAULTS["local"])
    if provider == "local":
        base_url = config.OLLAMA_API_URL
    elif provider == "vllm":
        base_url = config.VLLM_API_URL
    else:
        base_url = defaults["base_url"]
    c = LLMClient(
        base_url=base_url,
        api_key=api_key if api_key else "EMPTY",
        model=model,
        provider=provider,
    )
    # Yerel sağlayıcılarda sunucudan gerçek model adını al
    if provider in ("local", "vllm"):
        c.check_connection()
    return c

# Model adını seçili sağlayıcıya göre al
_selected_model = (
    config.OLLAMA_MODEL if st.session_state.llm_provider == "local"
    else config.VLLM_MODEL if st.session_state.llm_provider == "vllm"
    else PROVIDER_DEFAULTS.get(st.session_state.llm_provider, {}).get("model", "")
)

client = get_llm_client(
    st.session_state.llm_provider,
    st.session_state.get("llm_api_key", ""),
    _selected_model,
)

# --- SIDEBAR: LLM SAĞLAYICI SEÇİMİ ---
_PROVIDER_LABELS = {
    "local":  "🖥️ Local (Ollama)",
    "vllm":   "⚡ vLLM (Docker)",
    "openai": "🤖 ChatGPT (OpenAI)",
    "gemini": "✨ Gemini (Google)",
    "claude": "🧠 Claude (Anthropic)",
}

_LOCAL_PROVIDERS = {"local", "vllm"}

with st.sidebar:
    st.header("⚙️ LLM Ayarları")

    selected_provider = st.radio(
        "Sağlayıcı:",
        options=list(_PROVIDER_LABELS.keys()),
        format_func=lambda x: _PROVIDER_LABELS[x],
        key="llm_provider",
    )

    if selected_provider == "local":
        st.caption(f"Model: `{client.model}`")
        st.caption(f"URL: `{config.OLLAMA_API_URL}`")
        _api_key = ""
    elif selected_provider == "vllm":
        st.caption(f"Model: `{client.model}`")
        st.caption(f"URL: `{config.VLLM_API_URL}`")
        st.caption("Docker: `docker compose up -d`")
        _api_key = ""
    else:
        _placeholder = {
            "openai": "sk-...",
            "gemini": "AIza...",
            "claude": "sk-ant-...",
        }.get(selected_provider, "")
        _api_key = st.text_input(
            "API Anahtarı:",
            type="password",
            key="llm_api_key",
            placeholder=_placeholder,
        )
        _default_model = PROVIDER_DEFAULTS[selected_provider]["model"]
        st.caption(f"Varsayılan model: `{_default_model}`")
        if not _api_key:
            st.warning("API anahtarı girilmeden soru üretilemez.")



# --- DB CHECK ---
if not os.path.exists(DB_NAME):
    st.error("Veritabanı (maarif_gen.db) bulunamadı. Lütfen önce 'db_setup.py' ve 'etl_pipeline.py' betiklerini çalıştırın.")
    st.stop()

# --- MAIN CONTENT & SELECTIONS ---
st.title("🎓 Maarif-Gen: Otomatik Soru Üretim Sihirbazı")

col_head1, col_head2 = st.columns([4, 1])
with col_head2:
    if st.button("🔄 Verileri Yenile / Önbelleği Temizle", use_container_width=True):
        st.cache_data.clear()
        st.cache_resource.clear()
        for key in ['w_step', 's_subj', 's_grade', 's_unit',
                    'selected_outcome', 'selected_components',
                    'generated_context', 'generated_questions',
                    'context_elapsed', 'last_outcome_code']:
            if key in st.session_state:
                del st.session_state[key]
        st.rerun()

st.markdown("### 📌 Kriter Seçimi", help="Lütfen sırasıyla ders, sınıf, ünite/tema ve öğrenme çıktısı seçiniz.")

# Add slight styling for selectboxes to look more premium
st.markdown(
    """
    <style>
    .stSelectbox label {
        font-weight: 600 !important;
        color: #1E3A8A;
    }
    .stSelectbox div[data-baseweb="select"] {
        border-radius: 8px;
    }
    </style>
    """,
    unsafe_allow_html=True,
)

# 1. Select Subject
subjects = get_subjects()
if not subjects:
    st.error("Veritabanında ders bulunamadı.")
    st.stop()

def change_step(s):
    st.session_state.w_step = s

subject_names = {s['id']: s['name'] for s in subjects}

# --- STEP 1: DERS ---
if st.session_state.w_step == 1:
    st.markdown("#### 1. Ders Seçimi")
    s_id = st.selectbox("Ders Seçiniz:", options=list(subject_names.keys()), format_func=lambda x: subject_names[x])
    if st.button("Kaydet ve Devam Et ✨", key="btn_s1"):
        st.session_state.s_subj = s_id
        st.session_state.w_step = 2
        st.rerun()
elif st.session_state.w_step > 1:
    col1, col2 = st.columns([6, 1])
    col1.success(f"**📚 Ders:** {subject_names.get(st.session_state.s_subj)}")
    col2.button("✏️ Değiştir", key="edit_1", on_click=change_step, args=(1,), use_container_width=True)

# --- STEP 2: SINIF ---
if st.session_state.w_step == 2:
    st.markdown("#### 2. Sınıf Seçimi")
    grades = get_grades(st.session_state.s_subj)
    if not grades:
        st.warning("Bu derse ait sınıf bulunamadı.")
    else:
        grade_names = {g['id']: g['level_name'] for g in grades}
        g_id = st.selectbox("Sınıf Seçiniz:", options=list(grade_names.keys()), format_func=lambda x: grade_names[x])
        if st.button("Kaydet ve Devam Et ✨", key="btn_s2"):
            st.session_state.s_grade = g_id
            st.session_state.w_step = 3
            st.rerun()
elif st.session_state.w_step > 2:
    grades = get_grades(st.session_state.s_subj)
    grade_names = {g['id']: g['level_name'] for g in grades}
    col1, col2 = st.columns([6, 1])
    col1.success(f"**🏫 Sınıf:** {grade_names.get(st.session_state.s_grade, '')}")
    col2.button("✏️ Değiştir", key="edit_2", on_click=change_step, args=(2,), use_container_width=True)

# --- STEP 3: ÜNİTE ---
if st.session_state.w_step == 3:
    st.markdown("#### 3. Ünite/Tema Seçimi")
    units = get_units(st.session_state.s_grade)
    if not units:
        st.warning("Bu sınıfta ünite bulunamadı.")
    else:
        unit_names = {u['id']: u['name'] for u in units}
        u_id = st.selectbox("Ünite/Tema Seçiniz:", options=list(unit_names.keys()), format_func=lambda x: unit_names[x])
        if st.button("Kaydet ve Devam Et ✨", key="btn_s3"):
            st.session_state.s_unit = u_id
            st.session_state.w_step = 4
            st.rerun()
elif st.session_state.w_step > 3:
    units = get_units(st.session_state.s_grade)
    unit_names = {u['id']: u['name'] for u in units}
    col1, col2 = st.columns([6, 1])
    col1.success(f"**📖 Ünite/Tema:** {unit_names.get(st.session_state.s_unit, '')}")
    col2.button("✏️ Değiştir", key="edit_3", on_click=change_step, args=(3,), use_container_width=True)

# --- STEP 4: ÖĞRENME ÇIKTISI ---
if st.session_state.w_step == 4:
    st.markdown("#### 4. Öğrenme Çıktısı Seçimi")
    outcomes = get_outcomes(st.session_state.s_unit)
    if not outcomes:
        st.warning("Bu ünitede öğrenme çıktısı bulunamadı.")
    else:
        outcome_options = {o['id']: f"{o['code']} - {o['description'][:100]}..." if o['description'] else f"{o['code']} - (Metin Yok)" for o in outcomes}
        o_id = st.selectbox("Öğrenme Çıktısı Seçiniz:", options=list(outcome_options.keys()), format_func=lambda x: outcome_options[x])
        if st.button("Kaydet ve Başla 🚀", key="btn_s4", type="primary"):
            current_outcome = next(o for o in outcomes if o['id'] == o_id)
            current_components = get_components(current_outcome['id'])
            st.session_state.selected_outcome = current_outcome
            st.session_state.selected_components = current_components
            if st.session_state.last_outcome_code != current_outcome['code']:
                st.session_state.generated_context = ""
                st.session_state.generated_questions = {}
                st.session_state.context_elapsed = None
                st.session_state.last_outcome_code = current_outcome['code']
            st.session_state.w_step = 5
            st.rerun()
elif st.session_state.w_step > 4:
    col1, col2 = st.columns([6, 1])
    out_code = st.session_state.selected_outcome['code']
    out_desc = st.session_state.selected_outcome.get('description', '')
    col1.success(f"**🎯 Öğrenme Çıktısı:** {out_code} - {out_desc}")
    col2.button("✏️ Değiştir", key="edit_4", on_click=change_step, args=(4,), use_container_width=True)
    st.markdown("---")

if st.session_state.w_step < 5:
    st.info("👆 Lütfen yukarıdaki seçim adımlarını tamamlayın.")
    st.stop()
    
outcome = st.session_state.selected_outcome
components = st.session_state.selected_components

st.header("1. Bağlam (Context) Oluşturma")
outcome_desc = outcome.get('description', '') or f"{outcome['code']}"

# Show Implementation Guide if present to inform the user what the AI will use
impl_guide = outcome.get('implementation_guide', '')
if impl_guide:
    with st.expander(f"📖 {outcome['code']} - Uygulama Esaslarını Görüntüle", expanded=False):
        st.markdown(f"*{impl_guide}*")

with st.container():
    context_area = st.text_area(
        "Bağlam Senaryosu",
        value=st.session_state.generated_context,
        height=_text_height(st.session_state.generated_context),
        placeholder="Henüz bir bağlam oluşturulmadı. Aşağıdaki butona basın."
    )
    st.session_state.generated_context = context_area

    col_btn_ctx, col_time_ctx = st.columns([2, 3])
    with col_btn_ctx:
        if st.button("Yapay Zeka ile Bağlam Oluştur / Yenile", key="btn_baglan_olustur"):
            with st.spinner("Yapay zeka bağlam üretiyor..."):
                t0 = time.time()
                client_res = client.generate_context(
                    outcome={"code": outcome['code'], "text": outcome_desc},
                    impl_guide=impl_guide
                )
                elapsed = time.time() - t0
                if not client_res.startswith("Bağlam üretilirken hata"):
                    st.session_state.generated_context = client_res
                    st.session_state.context_elapsed = elapsed
                    st.rerun()
                else:
                    st.error(client_res)
    with col_time_ctx:
        if st.session_state.context_elapsed:
            st.caption(f"⏱ Bağlam {st.session_state.context_elapsed:.1f} saniyede üretildi")

st.header("2. Süreç Bileşenleri ve Soru Üretimi")

rubric_points = st.number_input("Rubrik Toplam Puanı (Örn: 10, 20)", value=10, step=5)

if not components:
    st.warning("Bu çıktı için alt süreç bileşenleri (a, b, c...) bulunamadı.")
else:
    n_total = len(components)
    n_done = sum(1 for c in components if st.session_state.generated_questions.get(c['code']))
    st.caption(f"Toplam {n_total} süreç bileşeni  •  {n_done} tamamlandı")

    # İlk işlenmemiş bileşenin indeksini bul (otomatik açmak için)
    first_unprocessed = next(
        (i for i, c in enumerate(components)
         if not st.session_state.generated_questions.get(c['code'])),
        None
    )

    for i, comp in enumerate(components):
        comp_code = comp['code']
        q_data = st.session_state.generated_questions.get(comp_code)
        is_skipped = isinstance(q_data, dict) and q_data.get('_skipped')
        has_question = bool(q_data) and not is_skipped

        is_current = (i == first_unprocessed)
        is_future = first_unprocessed is not None and i > first_unprocessed

        # Expander başlığı: duruma göre ikon + sıra numarası
        if is_skipped:
            exp_label = f"⏭  {i+1}/{n_total} — {comp_code}) {comp['description']}"
        elif has_question:
            exp_label = f"✅  {i+1}/{n_total} — {comp_code}) {comp['description']}"
        else:
            exp_label = f"📝  {i+1}/{n_total} — {comp_code}) {comp['description']}"

        # Sırası henüz gelmeyen bileşenler sadece liste olarak gösterilir.
        # Böylece sırası geldiğinde expander ilk kez expanded=True ile render
        # edilir ve Streamlit'in "initial value" kuralı doğru çalışır.
        if is_future:
            st.markdown(f"⬜&nbsp; **{i+1}/{n_total}** — {comp_code}) {comp['description']}")
            continue

        with st.expander(exp_label, expanded=is_current):
            col_gen, col_skip = st.columns([1, 1])
            with col_gen:
                btn_label = "🔄 Yeniden Üret" if has_question else "Soru Hazırla"
                if st.button(btn_label, key=f"gen_{outcome['code']}_{comp_code}", use_container_width=True):
                    if not st.session_state.generated_context:
                        st.warning("Lütfen önce yukarıdan bir bağlam oluşturun.")
                    else:
                        with st.spinner(f"{comp_code} için soru üretiliyor..."):
                            ctx = _build_context_prompt(st.session_state.generated_context, impl_guide, rubric_points)
                            t0 = time.time()
                            res = client.generate_question(
                                outcome={"code": outcome['code'], "text": f"{comp['description']} ({outcome_desc})"},
                                context=ctx
                            )
                            res['_elapsed'] = round(time.time() - t0, 1)
                            st.session_state.generated_questions[comp_code] = res
                            st.rerun()
            with col_skip:
                skip_label = "↩ Geçişi Geri Al" if is_skipped else "Geç"
                if st.button(skip_label, key=f"skip_{outcome['code']}_{comp_code}", use_container_width=True):
                    if is_skipped:
                        del st.session_state.generated_questions[comp_code]
                    else:
                        st.session_state.generated_questions[comp_code] = {"_skipped": True}
                    st.rerun()

            # Üretilen soruyu göster
            if is_skipped:
                st.caption("Bu bileşen geçildi.")
            elif has_question:
                c1, c2 = st.columns([2, 1])
                with c1:
                    new_q_text = st.text_area(
                        "Soru Metni",
                        value=q_data.get('question_text', ''),
                        key=f"q_text_{outcome['code']}_{comp_code}",
                        height=_text_height(q_data.get('question_text', ''))
                    )
                    q_data['question_text'] = new_q_text

                    st.markdown("**Rubrik (Düzenlenebilir)**")
                    rubric_data = q_data.get('rubric', [])
                    if rubric_data:
                        rubric_df = pd.DataFrame(rubric_data)
                        edited_rubric = st.data_editor(
                            rubric_df,
                            num_rows="dynamic",
                            key=f"rubric_edit_{outcome['code']}_{comp_code}",
                            use_container_width=True
                        )
                        q_data['rubric'] = edited_rubric.to_dict('records')
                with c2:
                    if q_data.get('_elapsed'):
                        st.caption(f"⏱ {q_data['_elapsed']} saniyede üretildi")
                    st.info(f"**Bilişsel Düzey:** {q_data.get('cognitive_level', 'Belirtilmemiş')}")
                    st.info(f"**Doğru Cevap Özeti:**\n{q_data.get('correct_answer_summary', '')}")

# --- STEP 3: EXPORT ---
st.header("3. Dışa Aktar / Kaydet")

if st.session_state.generated_questions:
    subject_name = subject_names.get(st.session_state.s_subj, "Soru")
    
    st.info(f"Oluşturduğunuz soruları **{subject_name}.docx** ana soru bankası dosyasına ekleyebilirsiniz. Dosya yoksa sıfırdan oluşturulur, varsa bu sorular sonuna yapıştırılır.")
    
    if st.button(f"📥 {subject_name} Soru Bankasına Ekle (Kaydet)", type="primary"):
        saved_file = append_to_subject_docx(subject_name, outcome, components, st.session_state.generated_questions)
        
        # Ayrıca Fine-Tuning datasetine JSONL formatında ekle
        append_to_finetuning_dataset(
            client=client,
            subject_name=subject_name,
            outcome=outcome,
            components=components,
            questions=st.session_state.generated_questions,
            generated_context=st.session_state.generated_context,
            impl_guide=impl_guide,
            rubric_points=rubric_points
        )
        
        st.success(f"Sorular başarıyla '{saved_file}' ve 'finetune_dataset.jsonl' (.jsonl veri seti) dosyalarına EKLENDİ!")
        
    # Her ihtimale karşı güncel dosyayı indirme butonu
    current_file_path = os.path.join("SoruBankasi", f"{subject_name}.docx")
    if os.path.exists(current_file_path):
        with open(current_file_path, "rb") as f:
            st.download_button(
                label=f"💾 {subject_name} Soru Bankasını İndir (.docx)",
                data=f,
                file_name=f"{subject_name}_SoruBankasi.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
else:
    st.info("Soru ürettikten sonra kaydetme ve indirme butonları burada görünecektir.")
